from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
import os
import time
from time import sleep
from app_config import MULTILINGUAL_EMBEDDING_MODEL, EMBEDDING_MODEL, API_KEY, BASE_URL, LLM_MODEL
from app_config import SECONDARY_QUIZ_PROCESSING_PROMPT_TEMPLATE, MAIN_QUIZ_PROCESSING_PROMPT_TEMPLATE
from core.utils import clean_document_content, chunk_content,  purge_database,  ensure_sample_data_dir
from core.utils import  get_topic_focus_key, get_language_config, get_language_name, get_system_messages
import tempfile
from langchain_community.document_loaders import UnstructuredWordDocumentLoader, PyPDFLoader
from langchain.docstore.document import Document as LangChainDocument
from bertopic.backend import BaseEmbedder
import json
from datetime import datetime
import streamlit as st
import pypdf
from openai import OpenAI
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app_config import CHUNK_SIZE, CHUNK_OVERLAP
from sentence_transformers import SentenceTransformer
from sklearn.cluster import MiniBatchKMeans
from app_config import DESIRED_TOPICS
from bertopic import BERTopic
import random


# ==============================================
# Start Document Processor Functions
# ==============================================

def load_vector_db():
    
    """Load the FAISS vector database"""
    embeddings = HuggingFaceEmbeddings(model_name=MULTILINGUAL_EMBEDDING_MODEL)
    db_path = "sample_data/faiss_db"
    if os.path.exists(db_path):
        return FAISS.load_local(db_path, embeddings, allow_dangerous_deserialization=True)
    return None

def process_start_document(uploaded_file): 
    """Process uploaded document and save to FAISS database"""
    start_time = time.time()
    purge_database()
    db_path = "sample_data/faiss_db"
    file_extension = os.path.splitext(uploaded_file.name)[-1].lower()
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        if isinstance(uploaded_file, st.runtime.uploaded_file_manager.UploadedFile):
            temp_file.write(uploaded_file.read())
        else:
            temp_file.write(uploaded_file)
        temp_file_path = temp_file.name

    try:
        if file_extension == ".docx":
            loader = UnstructuredWordDocumentLoader(temp_file_path)
        elif file_extension == ".pdf":
            loader = PyPDFLoader(temp_file_path) #UnstructuredPDFLoader(temp_file_path)
        else:
            raise ValueError("Unsupported file format!")

        docs = loader.load()

        embeddings = HuggingFaceEmbeddings(model_name=MULTILINGUAL_EMBEDDING_MODEL)
        vector_db = FAISS.from_documents(docs, embeddings)

        ensure_sample_data_dir()       
        
        vector_db.save_local(db_path)
        
        file_path = os.path.join(db_path, uploaded_file.name)

        # Save file
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Create and save metadata
        metadata = {
            "document_name": uploaded_file.name,
            "file_size": uploaded_file.size,
            "file_type": file_extension,
            "embedding_model": str(embeddings.model_name),
            "document_pages": len(docs),
            "faiss_index_info": {
                "index_type": str(vector_db.index),
                "docstore_size": len(vector_db.docstore._dict),
                "index_size": os.path.getsize(os.path.join(db_path, "index.faiss"))
            },
            "created_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
        }

        # Save metadata to JSON file
        metadata_path = os.path.join(db_path, "metadata.json")
        with open(metadata_path, "w") as f:
            json.dump(metadata, f, indent=2)
        end_time = time.time()
        print(f"Document processed and loaded to FAISS db in {end_time - start_time:.2f} seconds")
        # Update session state to reflect new database status
        st.session_state["db_exists"] = True  
        st.session_state["document_processed"] = True  
        return vector_db
    finally:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
            
# ==============================================
# Topic Extractor Functions
# ==============================================

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file (PDF, DOCX, or TXT)"""
    file_extension = os.path.splitext(uploaded_file.name)[-1].lower()
    if file_extension == ".pdf":
        reader = pypdf.PdfReader(uploaded_file)
        return "\n".join([page.extract_text() for page in reader.pages])
    elif file_extension == ".docx":
        doc = Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    else:  # Assume plain text
        return uploaded_file.getvalue().decode("utf-8")
    
def chunk_content_tokens(content, max_tokens=8000, overlap=200):
    """Split content into chunks that fit within token limits"""
    chars_per_token = 4
    max_chars = max_tokens * chars_per_token
    
    # If content is small enough, return as is
    if len(content) <= max_chars:
        return [content]
    
    chunks = []
    start = 0
    
    while start < len(content):
        end = start + max_chars
        
        # Adjust end to not cut words
        if end < len(content):
            # Find the last period or newline before max_chars
            last_period = content.rfind(".", start, end)
            last_newline = content.rfind("\n", start, end)
            
            # Use the latest of these positions
            if last_period > 0 and last_period > last_newline:
                end = last_period + 1
            elif last_newline > 0:
                end = last_newline + 1
        
        chunks.append(content[start:end])
        
        # Move start position with overlap
        start = end - overlap
    
    return chunks

def extract_topics_from_chunk(chunk, max_topics, language_code="en"):
    
    if not API_KEY:
        st.error("API_KEY not found in environment variables")
        st.stop()
    
    chunk = clean_document_content(chunk[:8000])

    if BASE_URL is None or not BASE_URL.strip():
        client = OpenAI(api_key=API_KEY)
    else:
        client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

    # Define language-specific instructions
    
    system_messages = {
            "en": "You are a helpful assistant that always returns valid JSON with topical analysis.",
            "de": "Du bist ein hilfreicher Assistent, der immer gültiges JSON mit thematischer Analyse zurückgibt",
            "pl": "Jesteś pomocnikiem, który zawsze zwraca poprawny JSON z analizą tematyczną.",
            "ro": "Ești un asistent util care întotdeauna returnează JSON valid cu analiza topică.",
            "bg": "Ти си полезен асистент, който винаги връща валиден JSON с тематичен анализ."
        }

    system_message = system_messages.get(language_code, system_messages["en"])
    
    LANGUAGE_PROMPT = {
                "en": 
                    f"""
                    Extract {max_topics} DISTINCT key topics from the MAIN CONTENT of this document.
                    Focus on different aspects of the substantive content and IGNORE:
                    - Document metadata (version numbers, dates, etc.)
                    - Copyright notices
                    - Trademark information
                    - Legal disclaimers
                    - Header/footer content
                    - Administrative information

                    Return exactly {max_topics} topics in Enlish language in this structure:
                    {{
                        "topics": [
                            {{
                                 "aspect": "aspect",
                                 "topic_name": "specific name",
                                 "description": "detailed explanation",
                                 "importance": 1-5
                            }}
                            ]
                    }}
                    Content: {chunk[:8000]}
                    """
                ,
                "de": f"""
                        Extraktion von {max_topics} VERSCHIEDENEN Schlüsselthemen aus dem HAUPTINHALT dieses Dokuments.
                        Konzentriere dich auf verschiedene Aspekte des inhaltlichen Inhalts und IGNORIERE:
                        - Dokumentmetadaten (Versionsnummern, Daten usw.)
                        - Urheberrechtshinweise
                        - Markenininformationen
                        - Rechtliche Haftungsausschlüsse
                        - Kopf-/Fußzeileninhalt
                        - Verwaltungsinformationen

                        Gib genau {max_topics} Themen in deutscher Sprache in dieser Struktur zurück:
                        {{
                            "topics": [
                                {{
                                    "aspect": "Aspekt",
                                    "topic_name": "spezifischer Name",
                                    "description": "detaillierte Erklärung",
                                    "importance": 1-5
                                }}
                            ]
                        }}
                        Inhalt: {chunk[:8000]}
                        """,
                "pl": 
                    f"""
                    Wyodrębnij {max_topics} ODRĘBNYCH kluczowych tematów z TREŚCI GŁÓWNEJ tego dokumentu.
                    Skoncentruj się na różnych aspektach treści merytorycznej i IGNORUJ:

                    Metadane dokumentu (numery wersji, daty itp.)
                    Zastrzeżenia praw autorskich
                    Informacje o znakach towarowych
                    Oświadczenia prawne
                    Treść nagłówka/stopki
                    Informacje administracyjne
                    Zwróć dokładnie {max_topics} tematów w języku polskim w tej strukturze:
                    {{
                        "topics": [
                                {{
                                     "aspect": "aspekt",
                                     "topic_name": "konkretna nazwa",
                                     "description": "szczegółowe wyjaśnienie",
                                     "importance": 1-5
                                }}
                            ]
                    }}
                    Treść: {chunk[:8000]}

                    """
                ,
                "ro": 
                   f"""
                    Extrage {max_topics} teme cheie DISTINCTE din CONȚINUTUL PRINCIPAL al acestui document.
                    Concentrează-te asupra diferitelor aspecte ale conținutului substanțial și IGNORĂ:

                    Metadatele documentului (numere de versiune, date etc.)
                    Avertismente privind drepturile de autor
                    Informații despre mărci comerciale
                    Declarații juridice
                    Conținutul antetului/subsolului
                    Informații administrative
                    Returnează exact {max_topics} teme în limba română în această structură:
                    {{
                        "topics": [
                                {{
                                     "aspect": "aspect",
                                     "topic_name": "nume specific",
                                     "description": "explicație detaliată",
                                     "importance": 1-5
                                }}
                            ]
                    }}
                    Conținut: {chunk[:8000]}
                    """            
                ,
                "bg":
                    f"""
                    Екстрахирай {max_topics} РАЗЛИЧНИ ключови теми от ОСНОВНОТО СЪДЪРЖАНИЕ на този документ.
                    Фокусирай се върху различни аспекти на същественото съдържание и ИГНОРИРАЙ:

                    Метаданни на документа (версии, дати и т.н.)
                    Забележки за авторски права
                    Информация за търговски марки
                    Правни откази от отговорност
                    Съдържание на хедър/футър
                    Административна информация
                    Върни точно {max_topics} теми на Българския език в тази структура:
                    {{
                    "topics": [
                                {{
                                    "aspect": "аспект",
                                    "topic_name": "конкретно име",
                                    "description": "подробно обяснение",
                                    "importance": 1-5
                                }}
                            ]        
                    }}
                    Съдържание: {chunk[:8000]}
                     """                                     
                
    }
            
    language_code_prompt = LANGUAGE_PROMPT.get(language_code, LANGUAGE_PROMPT["en"])

    prompt = language_code_prompt
    
    #print(f"Prompt: {prompt}")
    try:
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            temperature=0.3,
            max_tokens=8192
        )
        response_content = response.choices[0].message.content
        try:
            result = json.loads(response_content)
            # Validate response structure
            if not isinstance(result.get("topics"), list):
                raise ValueError("Response missing 'topics' array")
            # Ensure we have some topics
            if not result["topics"]:
                raise ValueError("Empty topics array in response")
            return result
        except json.JSONDecodeError as e:
            print(f"JSON decode error: {e}")
            return {"topics": []}

    except Exception as e:
        st.error(f"API Error: {str(e)}")
        print(f"API Error: {str(e)}")
        return {"topics": []}
    
def process_chunk(index, chunk, max_topics, language_code="en"):
    print(f"Chunk {index + 1} language {language_code}")
    start = time.time()
    result = extract_topics_from_chunk(chunk, max_topics, language_code)
    duration = time.time() - start
    print(f"Chunk {index + 1} processed in {duration:.2f} seconds")
    sleep(1)
    return result
    
def process_document(document_text, document_name, max_topics=10, max_tokens=10000, language_code="en"):
    """Main processing function for uploaded documents"""
    print(f"Selected language: {language_code}")
    chunks = chunk_content(document_text, max_tokens) 
    print(f"Original content length: {len(document_text)}")
    print(f"Number of chunks: {len(chunks)}")

    all_topics = []
    start_time = time.time()
    results = parallel_topic_extraction(chunks, max_topics, language_code)
    end_time = time.time()
    print(f"All chunks processed in {end_time - start_time:.2f} seconds")
    for result in results:
        all_topics.extend(result.get("topics", []))  # Each result is a dict with "topics" key

    # Deduplicate and select top topics
    unique_topics = {}
    for topic in all_topics:
        name = topic.get("topic_name")
        if name and name not in unique_topics:
            unique_topics[name] = topic

    final_topics = list(unique_topics.values())[:max_topics]

    # Save results
    topics_data = {
        "document_name": document_name,
        "generated_at": datetime.now().isoformat(),
        "model_used": "deepseek-chat" ,
        "topics": final_topics,
        "language": language_code
    }
    os.makedirs("topics_data", exist_ok=True)
    topic_file = os.path.join("topics_data", "topics.json")
    with open(topic_file, "w") as f:
        json.dump(topics_data, f, indent=2)

    return topics_data

def parallel_topic_extraction(chunks, max_topics=5, language_code="en"):
    
    results = []
    with ThreadPoolExecutor(max_workers=5) as executor:  # Adjust based on API limits
        # Submit all tasks and keep track of their original indices
        futures = {
            executor.submit(process_chunk, i, chunk, max_topics, language_code): i
            for i, chunk in enumerate(chunks)
        }

        # Collect results and sort by original index
        results = [None] * len(chunks)  # Pre-allocate list
        for future in as_completed(futures):
            index = futures[future]
            results[index] = future.result()
    return results

# ==============================================
# Semantic Preprocessing And Topic Extraction
# ==============================================

def load_document(text):
    text =  clean_document_content(text)
    return [LangChainDocument(page_content=text)]

def chunk_text(docs):

    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_documents(docs)
    # Add position metadata to each chunk
    current_pos = 0
    for chunk in chunks:
        chunk.metadata['start_pos'] = current_pos
        current_pos += len(chunk.page_content)
    return chunks

def extract_topics(chunks, max_topics = DESIRED_TOPICS):

    texts = [chunk.page_content for chunk in chunks]

    # Edge case: Only 1 chunk → return it as a single topic
    if len(texts) == 1:
        return None, [0], texts  # Single topic

    # Edge case: 2 chunks → force 2 topics (KMeans needs n_clusters ≤ n_samples)
    if len(texts) == 2:
        model = SentenceTransformer(EMBEDDING_MODEL)
        embeddings = model.encode(texts)
        kmeans = MiniBatchKMeans(n_clusters=2)  # MiniBatchKMeans works for tiny data
        topics = kmeans.fit_predict(embeddings)
        return None, topics, texts  # Skip BERTopic for tiny docs

    # Default case: Use BERTopic with dynamic n_clusters
    model = SentenceTransformer(EMBEDDING_MODEL)
    embeddings = model.encode(texts, show_progress_bar=True)
    n_clusters = min(max_topics, len(texts) - 1)  # Ensure n_clusters < n_samples
    kmeans = MiniBatchKMeans(n_clusters=n_clusters)

    topic_model = BERTopic(
        embedding_model=DummyEmbedder(embeddings),
        hdbscan_model=kmeans,
        min_topic_size=1
    )
    topics, _ = topic_model.fit_transform(texts, embeddings)
    return topic_model, topics, texts

class DummyEmbedder(BaseEmbedder):
    def __init__(self, embeddings):
        self.embeddings = embeddings

    def embed(self, docs, verbose=False):
        return self.embeddings
    
def order_topics_by_position(texts, topics, chunks):
    """Return topic IDs ordered by their first appearance in document"""
    topic_positions = {}

    # Find first occurrence position for each topic
    for topic_id in set(topics):
        # Get indices of all chunks belonging to this topic
        chunk_indices = [i for i, t in enumerate(topics) if t == topic_id]

        # Find the earliest chunk position for this topic
        first_pos = min([chunks[i].metadata['start_pos'] for i in chunk_indices])
        topic_positions[topic_id] = first_pos

    # Sort topics by their first occurrence position
    return sorted(topic_positions.keys(), key=lambda x: topic_positions[x])

def summarize_topics(topic_model, topics, texts, language_code="en"):
    # API Setup
    if not API_KEY:
        raise ValueError("DEEPSEEK_API_KEY not found")
    
    if BASE_URL is None or not BASE_URL.strip():
        client = OpenAI(api_key=API_KEY)
    else:
        client = OpenAI(api_key=API_KEY, base_url=BASE_URL)

    topic_info = topic_model.get_topic_info()
    summarized_labels = {}
    
    system_messages = {
            "en": "You are a helpful assistant that always returns valid JSON with topical analysis.",
            "pl": "Jesteś pomocnikiem, który zawsze zwraca poprawny JSON z analizą tematyczną.",
            "ro": "Ești un asistent util care întotdeauna returnează JSON valid cu analiza topică.",
            "bg": "Ти си полезен асистент, който винаги връща валиден JSON с тематичен анализ."
        }

    system_message = system_messages.get(language_code, system_messages["en"])
    
    # Language-specific 
    response_language = {
            "en": "English",
            "pl": "Polish",
            "ro": "Romanian",
            "bg": "Bulgarian",
        }.get(language_code, "English")

    order = 1        
    for topic_id in topic_info['Topic']:
        topic_docs = [text for i, text in enumerate(texts) if topics[i] == topic_id][:5]
        LANGUAGE_PROMPT = {
                "en": 
                    f"""
                        The following texts are belong to the same topic cluster:

                        {chr(10).join(topic_docs)}

                        Summarize this cluster in 1 sentence and give a descriptive label.
                        Respond in {response_language} with this JSON structure:
                        {{
                            "topics": [
                                {{
                                    "topic_name": "specific name",
                                    "description": "detailed explanation"
                                }}
                            ]
                        }}

                    """
                ,
                "pl": 
                    f"""
                    Następujące teksty są należą do tego samego tematycznego klastra:

                    {chr(10).join(topic_docs)}

                    Podsumuj ten klastur w jednym zdaniu i podaj opisowy etykietę.
                    Odpowiedz w języku {response_language} za pomocą poniższej struktury JSON:
                    {{
                        "topics":[
                            {{
                                "topic_name": "konkretna nazwa",
                                "description": "szczegółowe wyjaśnienie"
                            }}
                        ]
                    }}
                """
                ,
                "ro": 
                   f"""
                    Următoarele texte și aparțin aceluiași grup tematic:

                    {chr(10).join(topic_docs)}

                    Rezumă acest grup într-o singură propoziție și dă un etichet descriptiv.
                    Răspunde în {response_language} cu următoarea structură JSON:
                    {{
                        "topics":[
                            {{
                                "topic_name": "nume specific",
                                "description": "explicație detaliată"
                            }}
                        ]
                    }}
                   """
                ,
                "bg": 
                    f"""
                    Следващите текстове са принадлежат към един и същ тематичен клъстер:

                    {chr(10).join(topic_docs)}

                    Обобщи този клъстер в едно изречение и дай описателен етикет.
                    Отговори на {response_language} със следната JSON структура:
                    {{
                        "topics": [
                            {{
                                "topic_name": "конкретно име",
                                "description": "подробно обяснение"
                            }}
                            ]
                    }}
                     """                                     
                
        }
        config = LANGUAGE_PROMPT.get(language_code, LANGUAGE_PROMPT["en"])
        if topic_model is None:
            prompt = f"""
            The following texts are written in {response_language} and are from a short document:

            {chr(10).join(texts)}

            Summarize the main topic in 1 sentence and give a descriptive label.
            Respond in {response_language} with this JSON structure:
            {{
                "topics": [
                    {{
                        "topic_name": "specific name",
                        "description": "detailed explanation"
                    }}
                ]
            }}
        """
        else:
            prompt = config
            #print(f"Prompt: {prompt}")
        try:
            print(f"Processing chunk {order}")
            start_time = time.time()
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"},
                temperature=0.3
            )
            #print(f"Raw response: {response}")
            # Safely parse JSON
            response_content = response.choices[0].message.content
            end_time = time.time()
            print(f"Chunk {order} processed in {end_time - start_time:.2f} seconds")
            order+=1
            #print(f"Content: {response_content}")
            result = json.loads(response_content)
            #print(f"JSON content: {result}")

            # Validate structure
            if not isinstance(result.get("topics"), list):
                raise ValueError(f"Invalid topics format: {result}")
            
        

            summarized_labels[topic_id] = result["topics"][0]  # Store first topic

        except Exception as e:
            print(f"Error processing topic {topic_id}: {str(e)}")
            summarized_labels[topic_id] = {
                "topic_name": f"Topic {topic_id} (Error)",
                "description": str(e)
            }

    return summarized_labels

def process_document_with_semantic_preprocessing(document_text, document_name, max_topics, language_code="en"):
    print("Loading document...")
    print(document_name)
    docs = load_document(document_text)
    print("Splitting into chunks...")
    chunks = chunk_text(docs)
    print(f"chunks = {len(chunks)}")
    if not chunks:
        print("⚠️ Document is empty.")
        return
    elif len(chunks) < 10:
        #Not processing short documents
        print(f"Processing SHORT document in {language_code}, using LLM directly")
        return process_document(document_text, document_name, max_topics, language_code=language_code)
    print("Extracting topics...")
    topic_model, topics, texts = extract_topics(chunks, max_topics)
    
    # Get topics in document order
    ordered_topic_ids = order_topics_by_position(texts, topics, chunks)
    
    print("Summarizing topics with LLM...")
    start_time = time.time()
    summaries = summarize_topics(topic_model, topics, texts, language_code=language_code)
    end_time = time.time()
    print(f"Content summarized in {end_time - start_time:.2f} seconds") 
    #print(f"Raw summaries: {summaries}")
    all_topics = []
    order = 1
    for topic_id in ordered_topic_ids:
        sr = summaries[topic_id]
        all_topics.append(sr)
        order+=1
    # Deduplicate and select top topics
    unique_topics = {}
    for topic in all_topics:
        name = topic.get("topic_name")
        if name and name not in unique_topics:
            unique_topics[name] = topic

    final_topics = list(unique_topics.values())[:max_topics]

    # Save results
    topics_data = {
        "document_name": document_name,
        "generated_at": datetime.now().isoformat(),
        "model_used": "deepseek-chat" ,
        "topics": final_topics,
        "language": language_code
    }
    os.makedirs("topics_data", exist_ok=True)
    topic_file = os.path.join("topics_data", "topics.json")
    with open(topic_file, "w") as f:
        json.dump(topics_data, f, indent=2)

    return topics_data

# ==============================================
# Quiz Generator Functions
# ==============================================

def query_faiss_for_quiz(vector_db):
    retriever = vector_db.as_retriever(search_type="similarity", search_kwargs={"k": 10})
    focus_topics = []
    if 'selected_topics' in st.session_state:
        focus_topics = st.session_state.selected_topics
  
    if focus_topics and len(focus_topics) > 0:
        topic_focus_prompt = f"focus specifically on these key topics: {', '.join(focus_topics)}"
    else:
        topic_focus_prompt = "focus on all available topics"
        
    results = retriever.invoke(f"Generate quiz based on stored documents, {topic_focus_prompt}")
    return " ".join([result.page_content for result in results]) if results else None

def generate_mcq_json(content, num_questions, difficulty_distribution, language_code="en"):

    if not API_KEY:
        st.error("API Key missing")
        return None
    
    if BASE_URL is None or not BASE_URL.strip():
        client = OpenAI(api_key=API_KEY)
    else:
        client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
    
    language_name = get_language_name(language_code)
    
    config = get_language_config(language_code)

    system_message = get_system_messages(language_code)

    # Difficulty prompt
    difficulty_instructions = []
    if difficulty_distribution["easy"] > 0:
            difficulty_instructions.append(
                f"{difficulty_distribution['easy']} {config['difficulty_levels'][0]} questions"
            )
    if difficulty_distribution["medium"] > 0:
            difficulty_instructions.append(
                f"{difficulty_distribution['medium']} {config['difficulty_levels'][1]} questions"
            )
    if difficulty_distribution["hard"] > 0:
        difficulty_instructions.append(
            f"{difficulty_distribution['hard']} {config['difficulty_levels'][2]} questions"
        )

    difficulty_prompt = "\n- ".join(difficulty_instructions)

    # Topic focus
    focus_topics = []
    if 'selected_topics' in st.session_state:
            focus_topics = st.session_state.selected_topics
            
    topic_focus_prompt = ""
    if focus_topics and len(focus_topics) > 0:
            topic_focus_key = get_topic_focus_key(language_code)
            topic_focus_prompt = f"\n- {topic_focus_key} {', '.join(focus_topics)}"
    
    print("Starting content chunking")
    start_time = time.time()
    content_chunks = chunk_content(content)
    random.shuffle(content_chunks)
    end_time = time.time()
    print(f"Content chunked in {end_time - start_time:.2f} seconds")
    generated_questions = set()
    all_questions = []
    processed_chunks = 0
    retry_attempts = 2  # Number of attempts to get unique questions per chunk
    fallback_attempts = 2  # Number of attempts to get fallback questions
    print(f"Content chunked in {len(content_chunks)} chunks")
    remaining_questions = num_questions
    # First pass - try to generate from content chunks
    for chunk in content_chunks:
            try:               
                print(f"Processing chunk {processed_chunks + 1}")            
                attempt = 0
                # Try to generate questions from this chunk
                while attempt <= retry_attempts:
                    print(f"Attempt {attempt+1}")
                    if remaining_questions <= 0:
                        break
                    random_seed = random.randint(0, 1000)
                    timestamp = int(time.time())
                    
                    prompt = MAIN_QUIZ_PROCESSING_PROMPT_TEMPLATE.format(
                            remaining_questions=remaining_questions,
                            language_name=language_name,
                            chunk=chunk,
                            random_seed=random_seed,
                            timestamp=timestamp,
                            create_different=config["instructions"]["create_different"],
                            options=config["instructions"]["options"],
                            randomize=config["instructions"]["randomize"],
                            difficulty_prompt=difficulty_prompt,
                            topic_focus_prompt=topic_focus_prompt,
                            difficulty=config["instructions"]["difficulty"],
                            avoid=config["instructions"]["avoid"],
                            plausible=config["instructions"]["plausible"],
                            explanations=config["instructions"]["explanations"],
                            uniqueness=config["instructions"]["uniqueness"],
                        )
                    start_time = time.time()
                    print("Calling API")
                    #print(prompt)
                    response = client.chat.completions.create(
                        model=LLM_MODEL,
                        messages=[
                            {"role": "system", "content": system_message},
                            {"role": "user", "content": prompt}
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.7
                    )
                    end_time = time.time()
                    print(f"Chunk processed in {end_time - start_time:.2f} seconds")
                    result = json.loads(response.choices[0].message.content)
                    #if "questions" in result:
                    new_questions = []
                    for q in result["questions"]:
                        # Basic duplicate check
                        question_text = q["question"].lower().strip()
                        if question_text not in generated_questions:
                            generated_questions.add(question_text)
                            new_questions.append(q)

                    if len(new_questions) > 0:
                        all_questions.extend(new_questions)
                        print(f"Generated {len(new_questions)} questions")
                    
                    print(f"All questions {len(all_questions)}")    
                    remaining_questions = num_questions - len(all_questions)
                    print(f"Remaining questions {remaining_questions}")
                    retry_attempts = retry_attempts - 1
                    attempt =  attempt+1      
                    if not result.get("content_adequate", True):
                            break  

                processed_chunks += 1
                end_time = time.time()
                print(f"Processed chunk in {end_time - start_time:.2f} seconds")

            except Exception as e:
                print(f"Failed to process chunk {processed_chunks + 1}: {str(e)}")
                st.warning(f"Failed to process chunk {processed_chunks + 1}: {str(e)}")
                continue

    # Generate fallback questions if we didn't get enough from the chunks
    if len(all_questions) < num_questions:
            remaining = num_questions - len(all_questions)
            print(f"Generating {remaining} fallback questions to reach requested total")
            focus_topics_str = ', '.join(focus_topics) if focus_topics else 'General concepts'
            existing_questions_str = ', '.join(q['question'][:50] + '...' for q in all_questions)
            for attempt in range(fallback_attempts):
                prompt = SECONDARY_QUIZ_PROCESSING_PROMPT_TEMPLATE.format(
                        remaining=remaining,
                        language_name=language_name,
                        focus_topics=focus_topics_str,
                        create_different=config['instructions']['create_different'],
                        options=config['instructions']['options'],
                        randomize=config['instructions']['randomize'],
                        difficulty_prompt=difficulty_prompt,
                        difficulty=config['instructions']['difficulty'],
                        avoid=config['instructions']['avoid'],
                        plausible=config['instructions']['plausible'],
                        explanations=config['instructions']['explanations'],
                        uniqueness=config['instructions']['uniqueness'],
                        fallback=config['instructions']['fallback'],
                        existing_questions=existing_questions_str,
                    )

                try:
                    response = client.chat.completions.create(
                        model=LLM_MODEL,
                        messages=[
                            {"role": "system", "content": system_message},
                            {"role": "user", "content": prompt}
                        ],
                        response_format={"type": "json_object"},
                        temperature=0.7  
                    )

                    result = json.loads(response.choices[0].message.content)
                    if "questions" in result:
                        new_questions = []
                        for q in result["questions"]:
                            question_text = q["question"].lower().strip()
                            if question_text not in generated_questions:
                                generated_questions.add(question_text)
                                new_questions.append(q)

                        if len(new_questions) > 0:
                            all_questions.extend(new_questions)
                            remaining-=len(all_questions)
                            if remaining <= 0:
                                break

                except Exception as e:
                    st.warning(f"Fallback question generation attempt {attempt + 1} failed: {str(e)}")
                    continue

    # Final quality check and trim to exact count
    final_questions = []
    seen_questions = set()

    for q in all_questions:
            question_text = q["question"].lower().strip()
            if question_text not in seen_questions:
                seen_questions.add(question_text)
                final_questions.append(q)
    
    random.shuffle(final_questions)

    # If we're still short (shouldn't happen but just in case)
    if len(final_questions) < num_questions:
            st.warning(f"Only able to generate {len(final_questions)} out of requested {num_questions} questions")

    # Build final quiz
    quiz_data = {
            "quiz_title": f"{config['quiz_title']} ({difficulty_distribution['easy']} {config['difficulty_levels'][0]}/"
                        f"{difficulty_distribution['medium']} {config['difficulty_levels'][1]}/"
                        f"{difficulty_distribution['hard']} {config['difficulty_levels'][2]})",
            "questions": final_questions[:num_questions],
            "focused_topics": focus_topics if focus_topics else "All topics"
        }

    return quiz_data
  
            

